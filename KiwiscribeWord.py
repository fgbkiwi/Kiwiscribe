import os
import sys
import re
import glob
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Try to import PyPDF2 for PDF parsing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# Function to select file via PyQt6, Tkinter, or Console
def select_file_gui(title="Selecione o arquivo", filter_desc="Arquivos de texto (*.txt)"):
    # Try using PyQt6 if available
    try:
        from PyQt6.QtWidgets import QApplication, QFileDialog
        app = QApplication.instance()
        if not app:
            app = QApplication(sys.argv)
        file_path, _ = QFileDialog.getOpenFileName(None, title, "", filter_desc)
        if file_path:
            return file_path
    except Exception:
        pass
        
    # Try using tkinter as fallback
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        ext = filter_desc.split("(")[1].split(")")[0] if "(" in filter_desc else "*.*"
        file_path = filedialog.askopenfilename(title=title, filetypes=[(filter_desc.split("(")[0].strip(), ext)])
        if file_path:
            return file_path
    except Exception:
        pass
        
    # Fallback to console input
    file_path = input(f"{title} ({filter_desc}): ").strip()
    if (file_path.startswith('"') and file_path.endswith('"')) or \
       (file_path.startswith("'") and file_path.endswith("'")):
        file_path = file_path[1:-1]
    return file_path

# Número de processo no padrão CNJ (com ou sem o sufixo completo).
_PROC_NUMBER_RE = r'\d{4,7}-\d{2}(?:\.\d{4}\.\d\.\d{2}\.\d{4})?'

def sanitize_foro_text(raw_foro):
    """Limpa texto de foro removendo ruídos de classe/número processual."""
    if not raw_foro:
        return ""

    foro = re.sub(r'\s+', ' ', raw_foro).strip()
    foro = re.sub(r'\s*-\s*', '-', foro)

    # Corta sufixos comuns de narrativa da ata.
    foro = re.split(
        r'\s+(?:realizou|audi[êe]ncia|relativa|sob|ata|processo|autos?)\b',
        foro,
        flags=re.IGNORECASE,
    )[0].strip()

    # Remove rótulos textuais de classe/processo (ex.: "- Ação Trabalhista").
    foro = re.sub(
        r'[\s-]+(?:Ação\s+Trabalhista|Proc\.?|Processo)\b.*$',
        '',
        foro,
        flags=re.IGNORECASE,
    ).strip()

    # O PJe quebra a linha entre a identificação da vara/foro e o número do
    # processo; ao extrair o PDF essa quebra vira espaço e a sigla da classe
    # processual (ATOrd, ATSum, RTOrd, RTSum, ACum, CartPrec...) acaba colada
    # ao nome do foro. Corta qualquer sigla seguida do número do processo.
    foro = re.sub(
        rf'[\s-]+[A-Za-zÀ-ÿ]{{2,12}}\s+{_PROC_NUMBER_RE}.*$',
        '',
        foro,
        flags=re.IGNORECASE,
    ).strip()

    # Remove número de processo isolado (sem sigla) e tudo o que vier depois.
    foro = re.sub(rf'[\s-]*\b{_PROC_NUMBER_RE}\b.*$', '', foro).strip()

    # Mantém apenas caracteres válidos para nomes de cidades/foros.
    foro = re.sub(r"[^A-Za-zÀ-ÖØ-öø-ÿ0-9\s\-.'/()]", ' ', foro)
    foro = re.sub(r'\s+', ' ', foro).strip(' -.,')
    return foro

def is_valid_foro_text(foro):
    """Valida foro para evitar preenchimento com lixo textual."""
    if not foro:
        return False

    # Evita termos claramente processuais no valor final.
    forbidden_patterns = [
        r'\bAT(?:Ord|Sum)\b',
        r'\bRT(?:Ord|Sum)\b',
        r'\bACum\b',
        r'\bACPCiv\b',
        r'\bCartPrec\b',
        r'\bprocesso\b',
        r'\bautos?\b',
        rf'\b{_PROC_NUMBER_RE}\b',
    ]
    combined = re.compile('|'.join(forbidden_patterns), re.IGNORECASE)
    return combined.search(foro) is None

# Function to extract metadata from the hearing minutes PDF
def extract_metadata_from_ata(pdf_path):
    metadata = {
        'num_vara': None,
        'foro': None,
        'juiz': None,
        'reclamante': None,
        'reclamada': None
    }
    if not PyPDF2:
        print("PyPDF2 não está instalado. Pulando extração de metadados do PDF.")
        return metadata
    if not os.path.exists(pdf_path):
        print(f"Arquivo de ata não encontrado: {pdf_path}")
        return metadata
        
    try:
        print(f"Lendo arquivo da ata: {pdf_path}")
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception:
                    pass
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # Normalize spaces
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Extract Vara and Foro
        vara_match = re.search(r'(\d+)[ªa]\s*Vara do Trabalho de\s*([^,.\n]+)', text, re.IGNORECASE)
        if vara_match:
            metadata['num_vara'] = vara_match.group(1).strip()
            foro = sanitize_foro_text(vara_match.group(2).strip())
            metadata['foro'] = foro if is_valid_foro_text(foro) else None
            print(f"Vara encontrada: {metadata['num_vara']}, Foro: {metadata['foro']}")
            
        # Extract Judge
        juiz_pattern = r'Juiz\(?a?\)?(?: Federal)? do Trabalho\s*(?:Substituto\(?a?\)?\s*)?:?\s+([A-ZÀ-Úa-zà-ú\s\'\.-]+?)(?:,|\s(?:CPF|PRESIDENTE)|\s*$)';
        juiz_match = re.search(juiz_pattern, text, re.IGNORECASE)
        if juiz_match:
            metadata['juiz'] = ' '.join(juiz_match.group(1).strip().split()).title()
            print(f"Juiz encontrado: {metadata['juiz']}")
            
        # Extract Claimant
        recl_pattern1 = r'Presente a parte reclamante ([^,]+), pessoalmente, acompanhado\(?a?\)? de seu\(?a?\)? advogado\(?a?\)'
        reclamante_match = re.search(recl_pattern1, text, re.IGNORECASE)
        if reclamante_match:
            metadata['reclamante'] = ' '.join(reclamante_match.group(1).strip().split()).title()
        else:
            # Captura até o próximo rótulo do cabeçalho (ADVOGADO/RECLAMADO),
            # uma vírgula ou o fim do texto — evita truncar o nome em 1 letra.
            recl_pattern3 = r'RECLAMANTE\(?A?\)?\s*:\s*(.+?)\s*(?:,|\bADVOGAD[OA]\b|\bRECLAMAD[OA]\b|$)'
            reclamante_match = re.search(recl_pattern3, text, re.IGNORECASE)
            if reclamante_match:
                metadata['reclamante'] = ' '.join(reclamante_match.group(1).strip().split()).title()
        if metadata['reclamante']:
            print(f"Reclamante encontrado: {metadata['reclamante']}")
                
        # Extract Respondent
        reclamada_pattern1 = r'Presente a parte reclamada ([^,]+), representado\(?a?\)? pelo\(?a?\)? preposto\(?a?\)'
        match_reclamada = re.search(reclamada_pattern1, text, re.IGNORECASE)
        if match_reclamada:
            metadata['reclamada'] = ' '.join(match_reclamada.group(1).strip().split()).title()
        else:
            reclamada_pattern3 = r'RECLAMAD[OA]\(?A?\)?\s*:\s*(.+?)\s*(?:,|\bADVOGAD[OA]\b|\bRECLAMANTE\b|$)'
            match_reclamada = re.search(reclamada_pattern3, text, re.IGNORECASE)
            if match_reclamada:
                metadata['reclamada'] = ' '.join(match_reclamada.group(1).strip().split()).title()
        if metadata['reclamada']:
            print(f"Reclamada encontrada: {metadata['reclamada']}")
                
    except Exception as e:
        print(f"Erro ao extrair metadados do PDF da ata: {e}")
        
    return metadata

# Bookmark replacement function for single runs
def replace_bookmark_in_para(paragraph, bookmark_name, text):
    p_element = paragraph._p
    b_starts = p_element.findall(qn('w:bookmarkStart'))
    for b_start in b_starts:
        if b_start.get(qn('w:name')) == bookmark_name:
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = text
            r.append(t)
            idx = p_element.index(b_start)
            p_element.insert(idx + 1, r)
            return True
    return False

def insert_text_at_bookmark(doc, bookmark_name, text):
    # Search in main body paragraphs
    for paragraph in doc.paragraphs:
        if replace_bookmark_in_para(paragraph, bookmark_name, text):
            return True
            
    # Search in header paragraphs
    for section in doc.sections:
        for h_type in ['header', 'first_page_header', 'even_page_header']:
            header = getattr(section, h_type, None)
            if header:
                for paragraph in header.paragraphs:
                    if replace_bookmark_in_para(paragraph, bookmark_name, text):
                        return True
                        
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_bookmark_in_para(paragraph, bookmark_name, text):
                        return True
    return False

# Paragraph insertion logic
def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    if style:
        new_para.style = style
    return new_para

def insert_paragraphs_at_bookmark(doc, bookmark_name, text_lines):
    target_para = None
    for p in doc.paragraphs:
        for b in p._p.findall(qn('w:bookmarkStart')):
            if b.get(qn('w:name')) == bookmark_name:
                target_para = p
                break
        if target_para is not None:
            break

    if target_para is None:
        return False

    style = target_para.style
    current_para = target_para
    inserted = False

    for line in text_lines:
        line_str = line.strip()
        if not line_str:
            continue
        current_para = insert_paragraph_after(current_para, line_str, style)
        inserted = True

    # Só consome o parágrafo-marcador quando há conteúdo para colocar no lugar.
    # Para seções sem depoimento, preserva o marcador (DepRcte, DepRcdo,
    # DepTestRcte1, ...) sob o título, em vez de apagá-lo silenciosamente e
    # deixar o título órfão — o que dificultava recolocar o trecho depois.
    if inserted:
        p_element = target_para._p
        p_element.getparent().remove(p_element)
    return inserted

# Pure-Python transcript parser and segmenter
def simplify_role(role, name, text_context=""):
    role_lower = role.lower()
    name_lower = name.lower() if name else ""
    
    if "testemunha" in role_lower:
        return "Testemunha"
        
    if "reclamante" in role_lower:
        if "advogado" in role_lower or "advogada" in role_lower:
            if "amanda" in name_lower or "gabriela" in name_lower or "advogada" in role_lower:
                return "Advogada do Reclamante"
            return "Advogado do Reclamante"
        return "Reclamante"
        
    if "reclamada" in role_lower or "preposto" in role_lower or "preposta" in role_lower:
        if "advogado" in role_lower or "advogada" in role_lower:
            if "gabrielle" in name_lower or "advogada" in role_lower:
                return "Advogada da Reclamada"
            return "Advogado da Reclamada"
        if "kamila" in name_lower or "camila" in name_lower or "preposta" in role_lower or "senhora" in text_context.lower():
            return "Preposta"
        return "Preposto"
        
    if "juiz" in role_lower or "juíza" in role_lower or "juiza" in role_lower:
        if "juíza" in role_lower or "juiza" in role_lower:
            return "Juíza"
        return "Juiz"
        
    return role

def clean_intro_lines(section_lines):
    first_speech_idx = None
    for idx, line in enumerate(section_lines):
        if line.get('is_deponent'):
            first_speech_idx = idx
            break
            
    if first_speech_idx is None:
        return section_lines
        
    keep_keywords = [
        "me ouve", "consegue ouvir", "escuta", "escutando", "ouvindo", 
        "contradita", "contraditar", "suspeito", "suspeição", "impedimento", 
        "protesto", "indeferida", "indeferido", "deferido", "deferida", 
        "senhor", "senhora", "qual", "como se chama", "cpf", "documento",
        "endereço", "endereco", "carteira", "identidade", "rg", "nome completo",
        "estado civil", "chamar", "trazer", "entrar", "compromisso", "falso testemunho",
        "responder a verdade", "falar a verdade", "preposto", "preposta", "testemunha"
    ]
    
    cleaned_intro = []
    for line in section_lines[:first_speech_idx]:
        text_lower = line['text'].lower()
        if any(kw in text_lower for kw in keep_keywords):
            cleaned_intro.append(line)
            
    return cleaned_intro + section_lines[first_speech_idx:]

def replace_witness_qualification(witness_lines):
    warning_keywords = ["falso testemunho", "falar a verdade", "responder a verdade", "compromisso", "crime de falso"]
    personal_keywords = ["cpf", "endereço", "endereco", "carteira", "identidade", "rg", "nome completo", "estado civil"]
    
    warning_start_idx = None
    warning_end_idx = None
    
    for idx, line in enumerate(witness_lines):
        text_lower = line['text'].lower()
        role_lower = line['role'].lower()
        
        if 'juiz' in role_lower and any(kw in text_lower for kw in warning_keywords):
            if warning_start_idx is None:
                warning_start_idx = idx
                
        if warning_start_idx is not None and idx > warning_start_idx:
            if 'testemunha' in role_lower:
                warning_end_idx = idx
                break
                
    if warning_start_idx is None or warning_end_idx is None:
        return witness_lines
        
    qualification_start_idx = warning_start_idx
    for idx in range(warning_start_idx - 1, -1, -1):
        line = witness_lines[idx]
        text_lower = line['text'].lower()
        role_lower = line['role'].lower()
        
        if any(kw in text_lower for kw in personal_keywords):
            qualification_start_idx = idx
        elif idx < qualification_start_idx - 2:
            break
            
    qual_timestamp = witness_lines[qualification_start_idx]['time']
    qual_placeholder = {
        'time': qual_timestamp,
        'role': 'Nota',
        'name': '',
        'text': f"({witness_lines[warning_end_idx]['role'].lower()} qualificada, advertida e compromissada)",
        'raw_format': f"[{qual_timestamp}] (testemunha qualificada, advertida e compromissada)"
    }
    
    new_lines = witness_lines[:qualification_start_idx] + [qual_placeholder] + witness_lines[warning_end_idx + 1:]
    return new_lines

_ORDINAL_MAP = {'1': '1', 'primeira': '1', '2': '2', 'segunda': '2', '3': '3', 'terceira': '3'}

def _party_from_label(text_lower):
    """Identifica a parte (Rcte/Rcdo) a partir de sinônimos no texto, ou None.

    Reclamante: 'reclamante', 'autor'. Reclamado: 'reclamad', 'ré'/'réu',
    'empresa'. Retorna None quando nenhuma parte é mencionada.
    """
    if 'reclamante' in text_lower or re.search(r'\bautor(?:a|es)?\b', text_lower):
        return 'Rcte'
    if ('reclamad' in text_lower or 'empresa' in text_lower
            or re.search(r'\br[eé](?:us?)?\b', text_lower)):
        return 'Rcdo'
    return None

def _witness_ordinal(role_lower):
    """Extrai o número ordinal de uma testemunha do rótulo (1/2/3), ou None."""
    m = re.search(r'(\d+)\s*[ªaº°]?\s*testemunha', role_lower)
    if m:
        return _ORDINAL_MAP.get(m.group(1))
    for word, value in _ORDINAL_MAP.items():
        if word.isalpha() and word in role_lower:
            return value
    return None

def classify_deponent_role(role):
    """Classifica o rótulo de uma fala quanto ao depoente.

    Retorna ('test', 'Rcte'|'Rcdo'|None, num|None) para testemunhas, ('dep',
    'Rcte'|'Rcdo', None) para reclamante/preposto/representante (a própria parte),
    ou None quando a fala não é de um depoente (advogado, juiz, interlocutor não
    identificado).

    Aceita rótulos com ordinal ("1ª Testemunha do Reclamante"), sem ordinal
    ("Testemunha da Reclamada") e sinônimos de parte ("Testemunha do Autor",
    "1ª Testemunha da Ré"). Quando o rótulo é apenas "Testemunha" (ou só o nome
    da testemunha), ``party``/``num`` vêm ``None`` e o chamador resolve a parte
    e a numeração pelo contexto e pela ordem das oitivas.
    """
    role_lower = role.lower()
    if 'advogado' in role_lower or 'advogada' in role_lower:
        return None
    if 'testemunha' in role_lower:
        return ('test', _party_from_label(role_lower), _witness_ordinal(role_lower))
    if ('representante' in role_lower and 'reclamad' in role_lower) \
            or 'preposto' in role_lower or 'preposta' in role_lower:
        return ('dep', 'Rcdo', None)
    if 'reclamante' in role_lower or re.fullmatch(r'\s*autor(?:a)?\s*', role_lower):
        return ('dep', 'Rcte', None)
    if 'reclamad' in role_lower or re.fullmatch(r'\s*r[eé](?:us?)?\s*', role_lower):
        return ('dep', 'Rcdo', None)
    return None

def _witness_party_hint(text_lower):
    """Pista, no diálogo, de qual parte é a testemunha (ex.: 'testemunha do autor')."""
    m = re.search(r'testemunh\w*\s+(?:\w+\s+){0,2}?d[oae]s?\s+(\w+)', text_lower)
    return _party_from_label(m.group(1)) if m else None

def _witness_call_cue(text_lower):
    """Indica a chamada/oitiva de uma (nova) testemunha; retorna a parte ou None."""
    if not re.search(r'\b(?:chamar?|ouvir|trazer|adentrar?|entrar|pr[oó]xim\w+|'
                     r'primeir\w+|segund\w+|terceir\w+|vamos)\b', text_lower):
        return None
    return _witness_party_hint(text_lower)

def segment_transcript(transcript_content):
    lines = transcript_content.split('\n')
    parsed_lines = []
    
    # Aceita carimbos de tempo em HH:MM:SS (ex.: [00:01:24]) e em MM:SS
    # (ex.: [01:24]); alguns pós-processadores emitem apenas minutos:segundos,
    # o que antes fazia a segmentação ignorar todas as linhas e esvaziar as
    # seções de depoimento.
    _ts = r'\d{1,2}:\d{2}(?::\d{2})?'
    pattern_three = re.compile(rf'^\[({_ts})\]\s*([^:]+?)\s*:\s*([^:]+?)\s*:\s*(.*)$')
    pattern_two = re.compile(rf'^\[({_ts})\]\s*([^:]+?)\s*:\s*(.*)$')
    
    for idx, line in enumerate(lines):
        line_str = line.strip()
        if not line_str:
            continue
            
        m3 = pattern_three.match(line_str)
        if m3:
            time_str, role, name, text = m3.groups()
            parsed_lines.append({
                'line_idx': idx,
                'time': time_str,
                'role': role.strip(),
                'name': name.strip(),
                'text': text.strip(),
                'section': None
            })
            continue
            
        m2 = pattern_two.match(line_str)
        if m2:
            time_str, role, text = m2.groups()
            parsed_lines.append({
                'line_idx': idx,
                'time': time_str,
                'role': role.strip(),
                'name': '',
                'text': text.strip(),
                'section': None
            })
            
    section_keys = [
        'DepTestRcte1', 'DepTestRcte2', 'DepTestRcte3',
        'DepTestRcdo1', 'DepTestRcdo2', 'DepTestRcdo3',
        'Rcte', 'Rcdo',
    ]

    # 1. Direct label assignment.
    #
    # A atribuição é por PARTE (não pela ordem em que os depoimentos ocorrem),
    # então a ordem das oitivas pode variar livremente: testemunhas da reclamada
    # antes das do reclamante, preposto antes do reclamante, ou uma parte
    # reinquirida após as testemunhas — cada fala vai para a seção da sua parte.
    #
    # Testemunhas sem ordinal são numeradas pela ordem das oitivas (cada bloco
    # contíguo da mesma parte = próxima testemunha 1ª/2ª/3ª); interrupções de
    # advogado/juiz não encerram o bloco. Rótulos só com "Testemunha" (ou apenas
    # o nome) têm a parte inferida por pistas do diálogo e pela continuidade.
    witness_blocks = {'Rcte': 0, 'Rcdo': 0}  # nº de testemunhas já numeradas/parte
    prev_deponent = None      # (kind, party) da última fala de depoente
    pending_party = None      # parte sugerida pela última menção a "testemunha d..."
    new_call = False          # houve chamada de nova testemunha desde a última fala
    unresolved_witness = False
    for line in parsed_lines:
        text_lower = line['text'].lower()
        hint = _witness_party_hint(text_lower)
        if hint:
            pending_party = hint
        call = _witness_call_cue(text_lower)
        if call:
            pending_party = call
            new_call = True

        cls = classify_deponent_role(line['role'])
        if cls is None:
            continue
        kind, party, num = cls

        if kind == 'dep':
            line['section'] = party
            line['is_deponent'] = True
            prev_deponent = (kind, party)
            new_call = False
            continue

        # --- testemunha ---
        block_party = prev_deponent[1] if (prev_deponent and prev_deponent[0] == 'test') else None
        if party is None:
            # "Testemunha" sem parte: usa nova chamada, senão a testemunha em
            # curso, senão a última pista; em último caso, a ordem legal padrão.
            if new_call and pending_party:
                party = pending_party
            elif block_party:
                party = block_party
            elif pending_party:
                party = pending_party
            else:
                party = 'Rcte'
                unresolved_witness = True

        if num is not None:
            section_num = num                       # ordinal explícito: direto
            witness_blocks[party] = max(witness_blocks[party], int(num))
        else:
            # Nova testemunha quando muda a parte ou houve chamada de testemunha.
            if party != block_party or new_call:
                witness_blocks[party] = min(witness_blocks[party] + 1, 3)
            section_num = str(max(witness_blocks[party], 1))

        line['section'] = f'DepTest{party}{section_num}'
        line['is_deponent'] = True
        prev_deponent = ('test', party)
        new_call = False

    if unresolved_witness:
        print("Aviso: alguma testemunha não pôde ser atribuída a uma parte pelo "
              "rótulo; usando a ordem padrão (reclamante). Confira o resultado.")

    # 2. Sequential/Transition filling.
    current_section = 'Rcte'
    for idx, line in enumerate(parsed_lines):
        if line['section'] is not None:
            current_section = line['section']
            continue

        text_lower = line['text'].lower()

        # Lookahead: próxima fala já rotulada como de um depoente.
        next_dep = None
        for future_idx in range(idx, len(parsed_lines)):
            if parsed_lines[future_idx]['section'] is not None:
                next_dep = parsed_lines[future_idx]['section']
                break

        if next_dep and next_dep != current_section:
            if next_dep == 'Rcdo':
                if any(kw in text_lower for kw in ['preposta', 'preposto', 'representante', 'chamar', 'consegue ouvir', 'trazer']):
                    current_section = 'Rcdo'
            elif next_dep.startswith('DepTest'):
                if any(kw in text_lower for kw in ['testemunha', 'contraditar', 'contradita', 'chamar', 'ouvir', 'escutando', 'trazer', 'compromisso']):
                    current_section = next_dep

        line['section'] = current_section

    # 3. Grouping and filtering
    sections_content = {k: [] for k in section_keys}
    for line in parsed_lines:
        sections_content[line['section']].append(line)
        
    final_sections = {}
    rcte_name = ""
    rcdo_name = ""
    
    for dep, dep_lines in sections_content.items():
        has_speech = False
        for line in dep_lines:
            if line.get('is_deponent'):
                has_speech = True
                # Store speaker name if deponent
                if dep == 'Rcte' and not rcte_name:
                    rcte_name = line['name'].upper()
                elif dep == 'Rcdo' and not rcdo_name:
                    rcdo_name = line['name'].upper()
                break

        if not has_speech:
            final_sections[dep] = []
            continue

        # Clean intro lines
        dep_lines = clean_intro_lines(dep_lines)
        
        # Replace qualifications for witnesses
        if dep.startswith('DepTest'):
            dep_lines = replace_witness_qualification(dep_lines)
            
        # Format final lines
        formatted_lines = []
        for line in dep_lines:
            if 'raw_format' in line:
                formatted_lines.append(line['raw_format'])
            else:
                sim_role = simplify_role(line['role'], line['name'], line['text'])
                formatted_lines.append(f"[{line['time']}] {sim_role}: {line['text']}")
        final_sections[dep] = formatted_lines
        
    return final_sections, rcte_name, rcdo_name

# Main workflow function
def generate_word_document(transcript_path, ata_pdf_path=None, output_dir=None, logger=print):
    """Gera o arquivo DOCX de degravação a partir de um TXT transcrito."""
    if not transcript_path or not os.path.exists(transcript_path):
        raise FileNotFoundError("Arquivo de transcrição não selecionado ou inexistente.")

    log = logger if callable(logger) else (lambda *_args, **_kwargs: None)
    log(f"Arquivo selecionado: {transcript_path}")

    filename = os.path.basename(transcript_path)
    proc_no = filename[:25]
    log(f"Número do Processo extraído do nome do arquivo: {proc_no}")

    txt_dir = os.path.dirname(transcript_path)
    target_dir = output_dir if output_dir else txt_dir

    metadata = {
        'num_vara': None,
        'foro': None,
        'juiz': None,
        'reclamante': None,
        'reclamada': None
    }

    resolved_ata_pdf = ata_pdf_path
    if not resolved_ata_pdf:
        pdf_files = glob.glob(os.path.join(txt_dir, f"*{proc_no}*.pdf"))
        if not pdf_files:
            pdf_files = glob.glob(os.path.join(txt_dir, "*ata*.pdf"))
            if not pdf_files:
                all_pdfs = glob.glob(os.path.join(txt_dir, "*.pdf"))
                if len(all_pdfs) == 1:
                    pdf_files = all_pdfs
        if pdf_files:
            resolved_ata_pdf = pdf_files[0]

    if resolved_ata_pdf and os.path.exists(resolved_ata_pdf):
        log(f"Arquivo de ata correspondente encontrado: {resolved_ata_pdf}")
        metadata = extract_metadata_from_ata(resolved_ata_pdf)
    else:
        log("Aviso: Nenhum arquivo PDF de ata correspondente encontrado. Continuando sem metadados do PDF.")

    with open(transcript_path, 'r', encoding='utf-8') as f:
        transcript_content = f.read()

    log("Segmentando e formatando depoimentos localmente...")
    segmented, rcte_parsed_name, rcdo_parsed_name = segment_transcript(transcript_content)
    log("Segmentação concluída!")

    script_dir = os.path.dirname(os.path.abspath(__file__))
    # Nome ASCII é o oficial (evita problemas de codificação no instalador NSIS);
    # o nome antigo com acentos é mantido como fallback para setups existentes.
    template_names = ["Degravacao_script.docx", "Degravação_script.docx"]
    template_path = None
    for base_dir in (script_dir, "."):
        for name in template_names:
            candidate = os.path.join(base_dir, name)
            if os.path.exists(candidate):
                template_path = candidate
                break
        if template_path:
            break
    if not template_path:
        raise FileNotFoundError(
            f"Template Degravacao_script.docx não encontrado em {script_dir} ou no diretório atual."
        )

    log(f"Usando template: {template_path}")
    doc = docx.Document(template_path)

    insert_text_at_bookmark(doc, "NumProc", proc_no)

    if metadata['num_vara']:
        num_vara_clean = re.sub(r'[ªa]$', '', metadata['num_vara']).strip()
        insert_text_at_bookmark(doc, "NumVara0", num_vara_clean)
    if metadata['foro']:
        foro_clean = sanitize_foro_text(metadata['foro'])
        if is_valid_foro_text(foro_clean):
            insert_text_at_bookmark(doc, "Foro1", foro_clean)
        else:
            log("Aviso: Foro extraído da ata foi descartado por conter texto inválido.")

    rcte_final_name = metadata['reclamante'] or rcte_parsed_name or ""
    rcdo_final_name = metadata['reclamada'] or ""

    if rcte_final_name:
        insert_text_at_bookmark(doc, "Rcte", rcte_final_name.upper())
    if rcdo_final_name:
        insert_text_at_bookmark(doc, "Rcdo", rcdo_final_name.upper())

    testimonies_map = {
        'DepRcte': segmented.get('Rcte', []),
        'DepRcdo': segmented.get('Rcdo', []),
        'DepTestRcte1': segmented.get('DepTestRcte1', []),
        'DepTestRcte2': segmented.get('DepTestRcte2', []),
        'DepTestRcte3': segmented.get('DepTestRcte3', []),
        'DepTestRcdo1': segmented.get('DepTestRcdo1', []),
        'DepTestRcdo2': segmented.get('DepTestRcdo2', []),
        'DepTestRcdo3': segmented.get('DepTestRcdo3', [])
    }

    for bookmark, lines in testimonies_map.items():
        insert_paragraphs_at_bookmark(doc, bookmark, lines)

    raw_lines = [l.strip() for l in transcript_content.split('\n') if l.strip()]
    insert_paragraphs_at_bookmark(doc, "Rascunho", raw_lines)

    output_filename = f"{proc_no} degravação.docx"
    output_path = os.path.join(target_dir, output_filename)

    doc.save(output_path)
    log("Documento de degravação criado com sucesso!")
    log(f"Caminho do arquivo gerado: {output_path}")
    return output_path

def main():
    print("=== Kiwiscribe Word Post-Processor ===")

    if len(sys.argv) > 1:
        transcript_path = sys.argv[1]
    else:
        transcript_path = select_file_gui("Selecione o arquivo de transcrição formatado", "Arquivos de texto (*.txt)")

    try:
        generate_word_document(transcript_path, logger=print)
    except Exception as e:
        print(f"Erro: {e}")
        sys.exit(1)

if __name__ == '__main__':
    main()
